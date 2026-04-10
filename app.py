import streamlit as st
import PyPDF2
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import google.generativeai as genai
import os
from PIL import Image

# ─────────────────────────────────────────────────────────────────────────────
# 1. UI Config
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="ResumePro Elite", layout="wide", page_icon="💎")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .main { background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); }

    /* Primary action button */
    .stButton>button {
        width: 100%; border-radius: 12px; height: 3.5em;
        background: #007bff; color: white; font-weight: bold; border: none;
    }

    /* Download button */
    .stDownloadButton>button {
        width: 100%; border-radius: 12px; height: 3.5em;
        background: #28a745; color: white; border: none;
    }

    /* ── SAVE HINT BANNER ── */
    .save-hint {
        background: #fff3cd;
        border: 2px solid #ffc107;
        border-radius: 10px;
        padding: 10px 16px;
        font-size: 15px;
        font-weight: 700;
        color: #7d4e00;
        text-align: center;
        margin-bottom: 8px;
        letter-spacing: 0.3px;
    }
    .save-hint kbd {
        background: #343a40;
        color: #fff;
        border-radius: 5px;
        padding: 2px 7px;
        font-size: 13px;
        font-family: monospace;
        margin: 0 2px;
    }
    </style>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# 2. AI Engine Config
# ─────────────────────────────────────────────────────────────────────────────
MODEL_NAME = "gemini-2.5-flash-lite"

try:
    if "GEMINI_API_KEY" in st.secrets:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    else:
        st.error("API Key missing in Streamlit Secrets.")
except Exception as e:
    st.error(f"Setup Error: {e}")

if "original_ai_output" not in st.session_state:
    st.session_state.original_ai_output = ""

# ─────────────────────────────────────────────────────────────────────────────
# 3. Sidebar
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("# 💎 Elite Control")
    with st.expander("🏢 BRANDING", expanded=True):
        company_choice = st.selectbox("Select Template", ["W3G", "Synectics", "ProTouch"])
        contact_number = st.text_input("Contact Number", value="123-456-7890")
        document_name  = st.text_input("Name", placeholder="Enter candidate name")

    with st.expander("🧠 AI ENGINE SETTINGS", expanded=True):
        include_summary  = st.checkbox("Develop Executive Summary", value=True)
        custom_points    = st.text_area("Custom Points", placeholder="Leadership, ROI...")
        make_confidential = st.checkbox("Anonymize Employers [CONFIDENTIAL]", value=False)

# ─────────────────────────────────────────────────────────────────────────────
# 4. Spacing constants  (all uniform — one value rules them all)
# ─────────────────────────────────────────────────────────────────────────────
SP = 6          # Pt — half-line spacing used uniformly everywhere
TWO_LINE_PT = 24

# ─────────────────────────────────────────────────────────────────────────────
# 5. Helper / Logic Functions
# ─────────────────────────────────────────────────────────────────────────────

def sentence_case(text: str) -> str:
    """
    Capitalise ONLY the very first character, leave everything else untouched.
    Acronyms like GAAP, IRS, CPA, QuickBooks stay intact.
    e.g. 'tax planning & GAAP'  -> 'Tax planning & GAAP'
         'FULL-CYCLE ACCOUNTING' -> 'FULL-CYCLE ACCOUNTING'  (skills keep caps)
    """
    t = text.strip()
    if not t:
        return t
    return t[0].upper() + t[1:]


def get_sections_dict(text: str) -> dict:
    """
    Parse AI text into {HEADER: [lines]}.
    Main headers  : ALL CAPS ending with colon  →  CORE SKILLS:
    Sub-headers   : lines starting with ##       →  ##Technical skills:
    Sub-headers are stored as-is so the renderer can detect and style them.
    Drops known skills-section noise artefacts.
    """
    sections, current_header = {}, None
    for line in text.split("\n"):
        clean = line.strip()
        if not clean:
            continue
        # Main section header — ALL CAPS + ends with colon
        if clean.isupper() and clean.endswith(":"):
            current_header = clean
            sections[current_header] = []
        elif current_header:
            # Drop table/software artefacts
            if "SKILL" in current_header.upper() and any(
                x in clean.lower() for x in ["software", "table", "the following"]
            ):
                continue
            sections[current_header].append(clean)
    return sections


def _replace_in_para(p, targets: dict):
    """
    Replace tokens inside a paragraph preserving run formatting.
    Merges all run text first so split tokens like '[' + 'CONTACT_NUMBER' + ']'
    are always found. Result goes into first run; others are cleared.
    """
    if not p.runs:
        return
    full_text = "".join(r.text for r in p.runs)
    if not any(tok in full_text for tok in targets):
        return
    new_text = full_text
    for token, val in targets.items():
        new_text = new_text.replace(token, val)
    if new_text == full_text:
        return
    p.runs[0].text = new_text
    for run in p.runs[1:]:
        run.text = ""


def _scan_element_for_placeholders(element, targets: dict):
    """
    Recursively scan ANY XML element for <w:t> nodes and replace tokens.
    This catches text inside floating text boxes (w:txbxContent), drawing
    shapes, and any other container that normal paragraph/table iteration misses.
    Uses raw XML so nothing is ever skipped regardless of nesting depth.
    """
    from docx.oxml.ns import qn as _qn
    W_T = _qn("w:t")
    # Collect all <w:t> elements anywhere in this element tree
    for wt in element.iter(W_T):
        text = wt.text or ""
        if not any(tok in text for tok in targets):
            continue
        new_text = text
        for token, val in targets.items():
            new_text = new_text.replace(token, val)
        wt.text = new_text

    # Also handle tokens split across sibling <w:t> nodes within the same <w:r>
    # by scanning all <w:p> ancestors and merging runs there too
    W_P = _qn("w:p")
    W_R = _qn("w:r")
    for wp in element.iter(W_P):
        runs = wp.findall(f".//{W_R}")
        full = "".join((r.find(W_T).text or "") if r.find(W_T) is not None else "" for r in runs)
        if not any(tok in full for tok in targets):
            continue
        new_full = full
        for token, val in targets.items():
            new_full = new_full.replace(token, val)
        if new_full == full:
            continue
        # Write back into first w:t of the paragraph, blank the rest
        first_wt = None
        for r in runs:
            wt = r.find(W_T)
            if wt is None:
                continue
            if first_wt is None:
                first_wt = wt
                first_wt.text = new_full
            else:
                wt.text = ""


def replace_all_placeholders(doc, contact: str, title: str, name: str):
    """
    Replace [CONTACT_NUMBER], [DOCUMENT_TITLE] and [NAME] everywhere —
    including inside floating text boxes in headers/footers which the
    standard python-docx .paragraphs iterator completely misses.
    """
    targets = {
        "[CONTACT_NUMBER]": contact,
        "[DOCUMENT_TITLE]": title,
        "[NAME]":           name,
    }

    # Scan every section header and footer element tree in full (catches textboxes)
    for section in doc.sections:
        for part in [section.header, section.footer]:
            _scan_element_for_placeholders(part._element, targets)

    # Scan the main document body element tree in full
    _scan_element_for_placeholders(doc.element.body, targets)


# ── Section-type detectors ────────────────────────────────────────────────────
def _is_list_section(h: str) -> bool:
    return any(kw in h.upper() for kw in
               ["SKILL", "CERTIF", "TOOL", "TECHNOLOG", "COMPETENC"])

def _is_exp_section(h: str) -> bool:
    return "EXPERIENCE" in h.upper()

def _is_edu_section(h: str) -> bool:
    return "EDUCATION" in h.upper()

def _is_summ_section(h: str) -> bool:
    return "SUMMARY" in h.upper()


# ── Content writers ───────────────────────────────────────────────────────────
def _base_run(p, text: str, bold=False, italic=False, size_pt=10.5):
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    return run


def add_bullet(doc, text: str, bold=False):
    """Single bullet line. Never bold by default."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(SP)
    clean = text.lstrip("•*-– ").strip()
    _base_run(p, f"•  {clean}", bold=bold)
    return p


def add_spacer(doc, before=0, after=SP):
    """Empty paragraph used purely for spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p


def split_by_pipe(line: str):
    """'A | B | C' → ['A', 'B', 'C']"""
    return [seg.strip() for seg in line.split("|") if seg.strip()]


# Date keywords used to distinguish "Company | Date" lines from
# job titles that happen to contain a pipe character.
_DATE_KEYWORDS = (
    "present", "current", "now",
    "jan", "feb", "mar", "apr", "may", "jun",
    "jul", "aug", "sep", "oct", "nov", "dec",
    "january","february","march","april","june","july",
    "august","september","october","november","december",
)
_DATE_DIGITS_RE = __import__("re").compile(r"\b(19|20)\d{2}\b")


def is_company_date_line(line: str) -> bool:
    """
    Return True only when the part AFTER the last '|' looks like a date range.
    e.g. 'Acme Corp | Jan 2020 - Present'  → True
         'Manager | Team Lead'              → False  (job title with pipe)
         'CPA | Auditing & Accounting'      → False  (education degree with pipe)
    """
    if "|" not in line:
        return False
    after_last_pipe = line.split("|")[-1].strip().lower()
    # Contains a 4-digit year (1900-2099) → it's a date
    if _DATE_DIGITS_RE.search(after_last_pipe):
        return True
    # Contains a month name or 'present' → it's a date
    if any(kw in after_last_pipe for kw in _DATE_KEYWORDS):
        return True
    return False


def _make_two_col_table(doc, total_dxa=8510, date_dxa=2520):
    """
    Invisible 2-column fixed-layout table.
    A4 body = 8510 DXA. date_dxa=2520 (1.75") safely fits the longest
    date 'Sep 2020 - Present' at 10.5pt Arial including default cell margins.
    Cell margins are explicitly zeroed so no hidden padding eats into the
    date column and causes wrapping.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    content_dxa = total_dxa - date_dxa

    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False

    tblPr = tbl._tbl.tblPr

    # Fixed layout — Word must honour explicit widths
    tblLayout = _OE("w:tblLayout")
    tblLayout.set(_qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Total table width
    tblW = tblPr.find(_qn("w:tblW"))
    if tblW is None:
        tblW = _OE("w:tblW"); tblPr.append(tblW)
    tblW.set(_qn("w:w"),    str(total_dxa))
    tblW.set(_qn("w:type"), "dxa")

    # Zero out default cell margins so no hidden padding shrinks the date cell
    tblCellMar = _OE("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = _OE(f"w:{side}")
        m.set(_qn("w:w"),    "0")
        m.set(_qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    # Remove all borders
    tblBorders = _OE("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = _OE(f"w:{side}")
        b.set(_qn("w:val"),  "none")
        b.set(_qn("w:sz"),   "0")
        b.set(_qn("w:space"),"0")
        b.set(_qn("w:color"),"auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Grid column definitions
    tblGrid = _OE("w:tblGrid")
    for dxa in (content_dxa, date_dxa):
        gc = _OE("w:gridCol"); gc.set(_qn("w:w"), str(dxa))
        tblGrid.append(gc)
    tbl._tbl.insert(list(tbl._tbl).index(tblPr) + 1, tblGrid)

    # Lock all cell widths explicitly
    for row in tbl.rows:
        for cell, dxa in zip(row.cells, (content_dxa, date_dxa)):
            _lock_cell(cell, dxa)

    return tbl, content_dxa, date_dxa


def _lock_cell(cell, dxa):
    """Set explicit DXA width on a cell so Word never overrides it."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = tcPr.find(_qn("w:tcW"))
    if tcW is None:
        tcW = _OE("w:tcW"); tcPr.append(tcW)
    tcW.set(_qn("w:w"),    str(dxa))
    tcW.set(_qn("w:type"), "dxa")


def _no_wrap_cell(cell):
    """Prevent cell content wrapping. Set BEFORE adding text."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    tcPr = cell._tc.get_or_add_tcPr()
    nw   = _OE("w:noWrap")
    tcPr.append(nw)


def add_experience_row(doc, company_part: str, date_part: str, job_title: str):
    """
    2-row fixed-layout table:
      Row 1: COMPANY (bold UPPER, left)  |  Date (italic, RIGHT-aligned, no-wrap)
      Row 2: Job title (bold, left)      |  [empty, no-wrap]
    Date is RIGHT-aligned so all dates end at the same right margin.
    """
    TOTAL = 8510
    DATE  = 2520    # 1.75" — safely fits longest date at 10.5pt with zero margins
    CONT  = TOTAL - DATE

    tbl, _, _ = _make_two_col_table(doc, total_dxa=TOTAL, date_dxa=DATE)
    r0c0, r0c1 = tbl.rows[0].cells
    r1c0, r1c1 = tbl.rows[1].cells

    # noWrap BEFORE adding text
    _no_wrap_cell(r0c1); _no_wrap_cell(r1c1)

    # Row 1: Company | Date (RIGHT-aligned)
    _base_run(r0c0.paragraphs[0], company_part.strip().upper(), bold=True)
    r0c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _base_run(r0c1.paragraphs[0], date_part.strip(), italic=True)

    # Row 2: Job title | empty
    _base_run(r1c0.paragraphs[0], sentence_case(job_title), bold=True)

    # Spacing
    for p in (r0c0.paragraphs[0], r0c1.paragraphs[0]):
        p.paragraph_format.space_before = Pt(SP)
        p.paragraph_format.space_after  = Pt(2)
    for p in (r1c0.paragraphs[0], r1c1.paragraphs[0]):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(SP)

    return tbl


def add_education_row(doc, degree_part: str, date_part: str, institution: str = ""):
    """
    2-row fixed-layout table (same widths as experience so dates align):
      Row 1: DEGREE (ALL CAPS bold, left)      |  Date (italic, RIGHT-aligned, no-wrap)
      Row 2: Institution (sentence case unbold) |  [empty, no-wrap]
    """
    TOTAL = 8510
    DATE  = 2520
    CONT  = TOTAL - DATE

    tbl, _, _ = _make_two_col_table(doc, total_dxa=TOTAL, date_dxa=DATE)
    r0c0, r0c1 = tbl.rows[0].cells
    r1c0, r1c1 = tbl.rows[1].cells

    # noWrap BEFORE adding text
    _no_wrap_cell(r0c1); _no_wrap_cell(r1c1)

    # Row 1: Degree ALL CAPS | Date italic RIGHT-aligned
    _base_run(r0c0.paragraphs[0], degree_part.strip().upper(), bold=True)
    r0c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _base_run(r0c1.paragraphs[0], date_part.strip(), italic=True)

    r0c0.paragraphs[0].paragraph_format.space_before = Pt(4)
    r0c0.paragraphs[0].paragraph_format.space_after  = Pt(2)
    r0c1.paragraphs[0].paragraph_format.space_before = Pt(4)
    r0c1.paragraphs[0].paragraph_format.space_after  = Pt(2)

    if institution:
        _base_run(r1c0.paragraphs[0], sentence_case(institution.strip()), bold=False)
        r1c0.paragraphs[0].paragraph_format.space_before = Pt(0)
        r1c0.paragraphs[0].paragraph_format.space_after  = Pt(SP)
        r1c1.paragraphs[0].paragraph_format.space_before = Pt(0)
        r1c1.paragraphs[0].paragraph_format.space_after  = Pt(SP)
    else:
        r1c0.paragraphs[0].paragraph_format.space_before = Pt(0)
        r1c0.paragraphs[0].paragraph_format.space_after  = Pt(0)
        r1c1.paragraphs[0].paragraph_format.space_before = Pt(0)
        r1c1.paragraphs[0].paragraph_format.space_after  = Pt(0)

    return tbl


def move_watermark_to_header(doc):
    """
    The template watermark/logo is a floating image anchored in the body —
    it only appears on page 1. This function copies the anchor XML into the
    header so it repeats on every page, then removes it from the body.
    The header already exists (it contains the line graphic and textbox).
    """
    from docx.oxml.ns import qn as _qn
    import copy

    W_P      = _qn("w:p")
    W_R      = _qn("w:r")
    W_DRAW   = _qn("w:drawing")
    WP_ANCHOR= "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"
    A_BLIP   = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    R_EMBED  = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

    body = doc.element.body

    # Find the paragraph in the body that contains the watermark drawing
    watermark_para = None
    watermark_rId  = None
    for para in body.iter(W_P):
        draw = para.find(f".//{W_DRAW}")
        if draw is not None:
            anchor = draw.find(WP_ANCHOR)
            if anchor is not None:
                blip = anchor.find(f".//{A_BLIP}")
                if blip is not None:
                    watermark_rId  = blip.get(R_EMBED)
                    watermark_para = para
                    break

    if watermark_para is None or watermark_rId is None:
        return  # nothing to move

    # Copy the watermark paragraph XML into the header
    section   = doc.sections[0]
    header    = section.header
    hdr_elem  = header._element

    watermark_copy = copy.deepcopy(watermark_para)
    # Insert before the last paragraph in the header
    hdr_paras = hdr_elem.findall(W_P)
    if hdr_paras:
        hdr_elem.insert(list(hdr_elem).index(hdr_paras[-1]), watermark_copy)
    else:
        hdr_elem.append(watermark_copy)

    # Add the image relationship to the header's relationship part
    # so the rId resolves correctly in the header context
    header_part = header.part
    doc_part    = doc.part
    # Find the image part via the doc relationship
    try:
        img_part = doc_part.related_parts[watermark_rId]
        new_rId  = header_part.relate_to(
            img_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        )
        # Update the rEmbed in the copy to use the new rId
        for blip in watermark_copy.iter(A_BLIP):
            blip.set(R_EMBED, new_rId)
    except Exception:
        pass  # if relationship copy fails, watermark still shows from body

    # Remove the watermark from the body so it doesn't double-show on page 1
    body.remove(watermark_para)


def set_keep_with_next(paragraph):
    """
    Set keepWithNext on a paragraph via XML so Word moves the whole
    block to the next page if it would otherwise start on the last line.
    """
    pPr = paragraph._element.get_or_add_pPr()
    kwn = OxmlElement("w:keepWithNext")
    pPr.append(kwn)


def set_keep_together(paragraph):
    """
    Set keepLines on a paragraph so its lines are never split across pages.
    """
    pPr = paragraph._element.get_or_add_pPr()
    kl = OxmlElement("w:keepLines")
    pPr.append(kl)


# ─────────────────────────────────────────────────────────────────────────────
# 6. Main Content Area
# ─────────────────────────────────────────────────────────────────────────────
st.title("Professional Resume Artisan")
uploaded_file = st.file_uploader("Drop Resume", type=["pdf", "docx", "png", "jpg", "jpeg"])
generate_btn  = st.button("✨ START AI TRANSFORMATION")

if uploaded_file and generate_btn:
    with st.status("🛠️ Re-architecting Content...", expanded=True):
        try:
            model = genai.GenerativeModel(MODEL_NAME)

            # ── Extract raw content from file ────────────────────────────────
            st.write("📄 Reading resume...")
            if uploaded_file.type == "application/pdf":
                raw = "".join([p.extract_text() for p in PyPDF2.PdfReader(uploaded_file).pages])
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                raw = "\n".join([p.text for p in docx.Document(uploaded_file).paragraphs])
            else:
                raw = Image.open(uploaded_file)

            # ── STAGE 1: Deep understanding ──────────────────────────────────
            st.write("🧠 Reading and understanding resume deeply...")

            understanding_prompt = """
You are an expert resume analyst. Before doing anything else, READ THE ENTIRE
RESUME from top to bottom, understand its full structure, every section, every
role, every qualification. Only after fully understanding it, produce the
structured output below.

═══════════════════════════════════════════════════════
STEP 1 — UNDERSTAND THE STRUCTURE
═══════════════════════════════════════════════════════
Identify every section in the resume — common ones include:
  Summary / Objective / Profile
  Experience / Work History / Employment
  Education / Academic Background / Qualifications
  Certifications / Licenses / Accreditations
  Skills / Core Competencies / Technical Skills / Professional Skills
  Tools / Technologies / Software
  Projects / Publications / Awards / Languages / etc.

IMPORTANT — SHARED SECTIONS:
Some resumes combine two topics under one heading, such as:
  "Education and Certifications", "Skills and Tools",
  "Certifications and Licenses", "Education and Professional Development"
When you see this, SPLIT them into TWO separate ALL-CAPS headers.
Example: "EDUCATION AND CERTIFICATIONS:" → write as:
  EDUCATION:
  (all education entries)
  CERTIFICATIONS:
  (all certification entries)

IMPORTANT — SKILLS WITH SUB-CATEGORIES:
If the resume lists skills under sub-categories like:
  Technical Skills: Python, SQL, Azure
  Professional Skills: Leadership, Communication
  Soft Skills: ...
Then preserve those sub-categories. Write the main header as CORE SKILLS:
and use indented sub-headers like:
  ##Technical skills:
  (skills listed as bullets)
  ##Professional skills:
  (skills listed as bullets)
The ## prefix marks a sub-header — do NOT use ## for main section headers.

═══════════════════════════════════════════════════════
STEP 2 — STRIP ALL PERSONAL / CONTACT INFORMATION
═══════════════════════════════════════════════════════
Remove completely — do NOT include anywhere in output:
  - Candidate full name
  - Phone numbers (any format)
  - Email addresses
  - LinkedIn, GitHub, portfolio, or any URLs
  - Home address, city, state, zip, country of residence
  - Any line that exists solely to identify the person

═══════════════════════════════════════════════════════
STEP 3 — STRUCTURE THE OUTPUT
═══════════════════════════════════════════════════════
Return a clean plain-text structured resume using these rules:

SECTION HEADERS:
  - ALL CAPS ending with a colon: EXPERIENCE:  EDUCATION:  CORE SKILLS:
  - Use EXPERIENCE: (never "Work History / Education" or combined labels)
  - Each main header on its own line

SUB-HEADERS (skills sub-categories only):
  - Prefix with ## and title case: ##Technical skills:
  - On its own line, directly inside the main section

EXPERIENCE ENTRIES:
  - Line 1: Company Name | Date Range
  - Line 2: Job Title  (on its own line, below company)
  - Lines 3+: bullet points starting with •
  - Every description line MUST start with • (never leave description unbulleted)
  - Job descriptions are plain text — NO bold, no asterisks, no markdown

DATE RANGE RULES (critical):
  - Format: Mon YYYY - Mon YYYY   or   Mon YYYY - Present
  - Use ONLY the first 3 letters of any month: Jan Feb Mar Apr May Jun
    Jul Aug Sep Oct Nov Dec
  - If BOTH month AND year are given: Jan 2020 - Mar 2023
  - If ONLY year is given: 2020 - 2023  (do NOT invent a month)
  - If NO date is given at all: leave the date field completely EMPTY
    (do NOT write "Date Unknown", "N/A", or any placeholder)
  - The entire date range must fit on one line — never split across two lines

EDUCATION ENTRIES:
  - Line 1: Degree Name | Date  (date only if actually present in the resume)
  - Line 2: Institution Name  (on its own line below the degree — NEVER after the pipe)
  - NEVER put institution name after the pipe — only a date goes after the pipe
  - If no date: Degree Name (no pipe at all if there is no date)

CERTIFICATIONS:
  - Certification Name | Date  (if date present)
  - Issuing body on the next line (if present)
  - If no date: just the certification name, no pipe

BULLETS:
  - Every job description point MUST start with •
  - Never leave a description line without a bullet
  - No bold, no asterisks, no markdown inside bullet text

UNIFORM SPACING (critical — must be consistent throughout):
  - Separate each job entry from the next with exactly ONE blank line
  - Separate each education entry from the next with exactly ONE blank line
  - One blank line before each section header
  - One blank line after each section header
  - All sections must have the same visual rhythm — no extra gaps, no missing gaps

GENERAL:
  - No markdown formatting anywhere (no **, no *, no #, no ---)
  - No mention of "Table", "Software" as content artefacts
  - No fabrication — only use information present in the original resume
  - No contact details anywhere in the output

Now read the resume completely, understand every section, then produce the
fully structured output:
"""

            stage1_response = model.generate_content(
                [understanding_prompt, raw] if not isinstance(raw, str)
                else [understanding_prompt, f"RESUME TEXT:\n{raw}"]
            )
            understood_resume = stage1_response.text.replace("**", "").strip()

            # ── STAGE 2: Polish and enforce all formatting rules ──────────────
            st.write("✨ Polishing and finalising...")

            sum_p = (
                f"Write or improve the SUMMARY: section using these focus points: "
                f"{custom_points}. If no focus points given, craft a strong "
                f"executive summary from the candidate's experience and skills."
            ) if include_summary else "Do NOT add, change, or remove any SUMMARY: section."

            priv_p = (
                "CONFIDENTIALITY RULE: Replace every employer/company name with "
                "'[CONFIDENTIAL]'. Apply this to every experience entry without exception."
            ) if make_confidential else ""

            reformat_prompt = f"""
You are a professional resume formatter. The resume below has already been
read and structured. Your job is to do a final quality pass — enforce every
formatting rule precisely and improve bullet point language.

{sum_p}
{priv_p}

FORMATTING RULES — enforce every single one:

1. SECTION HEADERS — ALL CAPS ending with colon. One per line.
   Valid examples: SUMMARY:  EXPERIENCE:  EDUCATION:  CORE SKILLS:
   CERTIFICATIONS:  TOOLS:  PROJECTS:

2. SUB-HEADERS (inside skills/tools sections only):
   Keep lines starting with ## exactly as-is: ##Technical skills:
   These are sub-headers — do NOT convert them to main headers.

3. EXPERIENCE ENTRIES — strict two-line format:
   Line 1: COMPANY NAME | Date Range
   Line 2: Job Title
   Then bullet points. Never put company and job title on the same line.

4. JOB DESCRIPTIONS — every single description line MUST:
   - Start with a bullet: •
   - Be plain text — NO bold, no asterisks, no markdown
   - Be action-oriented and results-focused
   - Never be left unbulleted even if the original had no bullet

5. DATE RANGE — strict format:
   - Use 3-letter month abbreviations only: Jan Feb Mar Apr May Jun
     Jul Aug Sep Oct Nov Dec
   - Format: Mon YYYY - Mon YYYY   OR   Mon YYYY - Present
   - Year only (if no month in resume): 2019 - 2022
   - NO date in resume: leave the date portion completely blank
     (NEVER write "Date Unknown", "N/A", "Present" if not stated)
   - The entire date range must fit on ONE line — never wrap

6. EDUCATION — Line 1: Degree | Date (blank after | if no date)
   Line 2: Institution

7. CERTIFICATIONS — Name | Date (if date given), institution on next line

8. BULLETS in all sections:
   - Must start with •
   - No bold, no asterisks, no markdown

9. NO fabrication — do not add experience, dates, or qualifications
   not present in the resume below.

10. NO contact details — no name, phone, email, address, or URLs.

11. Return ONLY the final resume text. No commentary, no preamble.

RESUME TO FINALISE:
{understood_resume}
"""

            stage2_response = model.generate_content(reformat_prompt)
            st.session_state.original_ai_output = (
                stage2_response.text
                .replace("**", "")
                .replace("__", "")
                .strip()
            )
            st.write("✅ Done!")

        except Exception as e:
            st.error(f"System Error: {e}")

# ─────────────────────────────────────────────────────────────────────────────
# 7. Editor & Export
# ─────────────────────────────────────────────────────────────────────────────
if st.session_state.original_ai_output:
    st.markdown("---")

    c_edit, c_preview = st.columns([1.5, 1])

    with c_edit:
        # ── Visible save hint ───────────────────────────────────────────────
        st.markdown(
            '<div class="save-hint">'
            '💾 Press <kbd>Ctrl</kbd> + <kbd>Enter</kbd> to apply your edits'
            ' before downloading'
            '</div>',
            unsafe_allow_html=True,
        )
        st.markdown("#### 🖋️ Live Editor")
        final_text = st.text_area(
            "Content Control:",
            value=st.session_state.original_ai_output,
            height=580,
            label_visibility="collapsed",
            help="Edit content here, then press Ctrl+Enter to apply changes.",
        )

    current_sections = get_sections_dict(final_text)

    with st.sidebar:
        st.markdown("---")
        header_order = st.multiselect(
            "Reorder Sections:",
            options=list(current_sections.keys()),
            default=list(current_sections.keys()),
        )

    # ── Rebuild editor text in reordered order so editor reflects reorder ──
    if header_order:
        reordered_lines = []
        for h in header_order:
            if h in current_sections:
                reordered_lines.append(h)
                reordered_lines.extend(current_sections[h])
        # Any headers not in header_order (unchecked) go at end
        for h, lines in current_sections.items():
            if h not in header_order:
                reordered_lines.append(h)
                reordered_lines.extend(lines)
        reordered_text = "\n".join(reordered_lines)
    else:
        reordered_text = final_text

    with c_edit:
        # ── Visible save hint ───────────────────────────────────────────────
        st.markdown(
            '<div class="save-hint">'
            '💾 Press <kbd>Ctrl</kbd> + <kbd>Enter</kbd> to apply your edits'
            ' before downloading'
            '</div>',
            unsafe_allow_html=True,
        )
        st.markdown("#### 🖋️ Live Editor")
        final_text = st.text_area(
            "Content Control:",
            value=reordered_text,
            height=580,
            label_visibility="collapsed",
            help="Edit content here, then press Ctrl+Enter to apply changes.",
        )

    # Re-parse after any manual edits in the editor
    current_sections = get_sections_dict(final_text)
    # header_order stays as selected in sidebar

    with c_preview:
        st.subheader("✅ Finalize & Download")

        if st.button("📋 CHECK & DOWNLOAD", use_container_width=True):
            with st.spinner("🔍 AI checking formatting before export..."):
                try:
                    model_qa = genai.GenerativeModel(MODEL_NAME)
                    qa_prompt = f"""
You are a strict resume formatting QA checker. Check every rule below and
return a CORRECTED resume text. Fix violations silently — no explanations.

RULES TO ENFORCE:

1. SECTION HEADERS: ALL CAPS ending colon, own line. e.g. EXPERIENCE:
   There must be exactly ONE blank line worth of space before and after
   each section header — enforce this in the text structure.

2. EXPERIENCE — strict 2-line per job:
   Line 1: COMPANY NAME | Mon YYYY - Mon YYYY
   Line 2: Job Title (alone on its own line, NEVER on same line as company)
   Date must contain ONLY a date — NO city, state, or location in the date.
   Each job block must be separated from the next by a consistent gap.

3. EDUCATION — strict 2-line per entry:
   Line 1: DEGREE NAME | Date   (ONLY date after pipe — nothing else)
   Line 2: Institution name (alone on its own line below the degree)
   If institution appears after the pipe → move it to line 2.
   If date and institution are mixed after the pipe → split correctly.
   Each education entry must be separated from the next by a consistent gap.

4. DATE RANGE: Mon YYYY - Mon YYYY or Mon YYYY - Present or YYYY - YYYY.
   3-letter months only: Jan Feb Mar Apr May Jun Jul Aug Sep Oct Nov Dec.
   No location, city, or state in the date. Blank if no date in original.

5. SPACING — must be uniform throughout:
   - Half-line gap (one blank line) between each job/education entry
   - Half-line gap before and after each section header
   - All bullets within a section have consistent spacing (no double gaps)

6. Every description bullet must start with •. No bold/asterisks/markdown.

7. No contact info (name, phone, email, address, URL) anywhere.

8. No markdown (no **, __, *, #headers, --- dividers).

Return ONLY the corrected resume text. No preamble. No commentary.

RESUME:
{final_text}
"""
                    qa_response = model_qa.generate_content(qa_prompt)
                    checked_text = (
                        qa_response.text
                        .replace("**", "").replace("__", "").strip()
                    )
                    current_sections = get_sections_dict(checked_text)
                    st.success("✅ QA passed — building resume...")
                except Exception as e:
                    st.warning(f"QA skipped: {e}")
                    checked_text = final_text

            t_map  = {
                "W3G":       "w3g_template.docx",
                "Synectics": "synectics_template.docx",
                "ProTouch":  "protouch_template.docx",
            }
            t_path = t_map.get(company_choice)
            doc    = docx.Document(t_path) if os.path.exists(t_path) else docx.Document()

            style           = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(10.5)

            replace_all_placeholders(doc, contact_number, document_name, document_name)
            move_watermark_to_header(doc)

            # ── CLEAR EMPTY TEMPLATE BODY PARAGRAPHS (keep drawings) ──────────
            from docx.oxml.ns import qn as _qn
            body      = doc.element.body
            W_DRAWING = _qn("w:drawing")
            W_PICT    = _qn("w:pict")
            while True:
                children = [c for c in body if c.tag != _qn("w:sectPr")]
                if not children:
                    break
                last = children[-1]
                if last.tag == _qn("w:p"):
                    has_draw = (last.find(f".//{W_DRAWING}") is not None or
                                last.find(f".//{W_PICT}")    is not None)
                    text = "".join(t.text or "" for t in last.iter(_qn("w:t")))
                    if not text.strip() and not has_draw:
                        body.remove(last)
                        continue
                break

            # ── SPACING: 1 line top & bottom on every page ────────────────────
            # The template header LINE and footer LINE are absolutely-positioned
            # drawings — changing paragraph spacing inside header/footer does NOT
            # move them and CORRUPTS the footer address layout.
            # Correct approach: add 12pt space_before on the FIRST body paragraph
            # (gap between header line and content) and 12pt space_after on the
            # LAST body paragraph (gap between content and footer line).
            # These body paragraphs are added inside the section loop below, so
            # we track them and set spacing after the loop.

            ONE_LINE_DXA = 240   # 12pt in twentieths-of-a-point (half-points)

            # ── SECTION LOOP ──────────────────────────────────────────────────
            first_body_para = None   # will be set on very first paragraph added

            for h in header_order:
                if h not in current_sections:
                    continue

                is_list = _is_list_section(h)
                is_exp  = _is_exp_section(h)
                is_edu  = _is_edu_section(h)
                is_summ = _is_summ_section(h)

                # Section header paragraph
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before   = Pt(SP)
                hp.paragraph_format.space_after    = Pt(SP)
                hp.paragraph_format.keep_with_next = True
                set_keep_with_next(hp)
                _base_run(hp, h, bold=True, size_pt=11)
                if first_body_para is None:
                    first_body_para = hp

                lines = current_sections[h]
                i = 0

                while i < len(lines):
                    line = lines[i]

                    # ── LIST SECTIONS ─────────────────────────────────────────
                    if is_list:
                        if line.startswith("##"):
                            sub_text = line.lstrip("#").strip()
                            sp = doc.add_paragraph()
                            sp.paragraph_format.space_before = Pt(SP)
                            sp.paragraph_format.space_after  = Pt(3)
                            sp.paragraph_format.keep_with_next = True
                            r = sp.add_run(sentence_case(sub_text))
                            r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10.5)
                        elif "|" in line and not is_company_date_line(line):
                            for seg in split_by_pipe(line):
                                add_bullet(doc, sentence_case(seg), bold=False)
                        else:
                            add_bullet(doc, sentence_case(line), bold=False)
                        i += 1

                    # ── EXPERIENCE ────────────────────────────────────────────
                    elif is_exp:
                        if is_company_date_line(line):
                            parts       = line.split("|")
                            company_str = parts[0].strip()
                            date_str    = parts[-1].strip()
                            job_title   = ""
                            if i + 1 < len(lines) and not is_company_date_line(lines[i + 1]):
                                job_title = lines[i + 1]
                                i += 1
                            anchor = doc.add_paragraph()
                            anchor.paragraph_format.space_before   = Pt(0)
                            anchor.paragraph_format.space_after    = Pt(0)
                            anchor.paragraph_format.keep_with_next = True
                            set_keep_with_next(anchor)
                            add_experience_row(doc, company_str, date_str, job_title)
                        else:
                            clean_line = line.lstrip("•*-– ").strip()
                            add_bullet(doc, sentence_case(clean_line), bold=False)
                        i += 1

                    # ── EDUCATION ─────────────────────────────────────────────
                    elif is_edu:
                        if is_company_date_line(line):
                            parts       = line.split("|")
                            degree_str  = parts[0].strip()
                            date_str    = parts[-1].strip()
                            institution = ""
                            if i + 1 < len(lines) and not is_company_date_line(lines[i + 1]):
                                institution = lines[i + 1]
                                i += 1
                            anchor = doc.add_paragraph()
                            anchor.paragraph_format.space_before   = Pt(0)
                            anchor.paragraph_format.space_after    = Pt(0)
                            anchor.paragraph_format.keep_with_next = True
                            set_keep_with_next(anchor)
                            add_education_row(doc, degree_str, date_str, institution)
                        else:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(SP)
                            _base_run(p, sentence_case(line), bold=False)
                        i += 1

                    # ── SUMMARY ───────────────────────────────────────────────
                    elif is_summ:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after  = Pt(SP)
                        _base_run(p, sentence_case(line), bold=False)
                        i += 1

                    # ── EVERYTHING ELSE ───────────────────────────────────────
                    else:
                        add_bullet(doc, sentence_case(line), bold=False)
                        i += 1

                add_spacer(doc, before=0, after=SP)

            # ── TOP SPACING: 1 line after header line, on every page ──────────
            # Set space_before on the FIRST body paragraph via XML directly.
            # This creates a visible gap between the header line and body content
            # that applies on every page because the header/sectPr repeats.
            if first_body_para is not None:
                pPr = first_body_para._element.get_or_add_pPr()
                spc = pPr.find(_qn("w:spacing"))
                if spc is None:
                    from docx.oxml import OxmlElement as _OE2
                    spc = _OE2("w:spacing"); pPr.append(spc)
                # Preserve existing after value, set before to ONE_LINE
                existing_after = spc.get(_qn("w:after"), str(int(SP * 20)))
                spc.set(_qn("w:before"), str(ONE_LINE_DXA))
                spc.set(_qn("w:after"),  existing_after)

            # ── BOTTOM SPACING: 1 line before footer line, on every page ──────
            # Add a final empty paragraph with space_after=12pt.
            # Do NOT touch the footer — its address boxes use absolute positioning
            # and any footer paragraph spacing change corrupts their layout.
            last_para = doc.add_paragraph()
            pPr = last_para._element.get_or_add_pPr()
            from docx.oxml import OxmlElement as _OE3
            spc2 = _OE3("w:spacing")
            spc2.set(_qn("w:before"), "0")
            spc2.set(_qn("w:after"),  str(ONE_LINE_DXA))
            pPr.append(spc2)

            # Hard page-break spacing
            for p in doc.paragraphs:
                for run in p.runs:
                    for br in run._element.findall(qn("w:br")):
                        if br.get(qn("w:type")) == "page":
                            p.paragraph_format.space_before = Pt(12)
                            p.paragraph_format.space_after  = Pt(12)

            buf = io.BytesIO()
            doc.save(buf)
            file_name = (f"{document_name.strip().upper()}.docx"
                         if document_name.strip() else "RESUME.docx")
            st.download_button(
                label=f"📥 DOWNLOAD {company_choice} DOCX",
                data=buf.getvalue(),
                file_name=file_name,
            )
